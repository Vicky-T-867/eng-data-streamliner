import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

SAFETY_LIMIT = 80


def analyze_data(excel_file):
    """Load sensor logs and pull rows over the safety limit."""
    df = pd.read_excel(excel_file)
    anomalies = df[df["Sensor_Value"] > SAFETY_LIMIT]
    avg_value = df["Sensor_Value"].mean()
    return df, anomalies, avg_value


def create_chart(df):
    """Simple time-series with the limit drawn as a dashed line."""
    plt.figure(figsize=(6, 4))
    plt.plot(df["Date"], df["Sensor_Value"], marker="o", color="b", label="Sensor Reading")
    plt.axhline(y=SAFETY_LIMIT, color="r", linestyle="--", label="Safety Limit")
    plt.title("Site Monitoring Trend")
    plt.xlabel("Date")
    plt.ylabel("Sensor Value")
    plt.legend()
    plt.tight_layout()

    chart_path = "temp_chart.png"
    plt.savefig(chart_path)
    plt.close()
    return chart_path


def generate_word_report(df, anomalies, avg_value, chart_path, output_docx):
    """Build a short status doc with summary, chart, and anomaly table."""
    doc = Document()
    doc.add_heading("Site Telemetry Monitoring Report", level=0)

    doc.add_heading("1. Summary", level=1)
    p = doc.add_paragraph("Average reading this period: ")
    p.add_run(f"{avg_value:.2f}").bold = True
    p.add_run(f". Threshold is {SAFETY_LIMIT}; pipeline flagged ")
    p.add_run(f"{len(anomalies)} anomaly event(s)").bold = True
    p.add_run(".")

    doc.add_heading("2. Trend Chart", level=1)
    doc.add_picture(chart_path, width=Inches(5.5))

    doc.add_heading("3. Anomaly Records", level=1)
    table = doc.add_table(rows=1, cols=3)
    headers = table.rows[0].cells
    headers[0].text = "Date"
    headers[1].text = "Location"
    headers[2].text = "Value"

    for _, row in anomalies.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["Date"])
        cells[1].text = str(row["Location"])
        cells[2].text = str(row["Sensor_Value"])

    doc.save(output_docx)
    print(f"Report saved: {output_docx}")


if __name__ == "__main__":
    target_data = "dummy_data.xlsx"
    output_report = "Project_Status_Report.docx"

    df, anomalies, avg_value = analyze_data(target_data)
    chart = create_chart(df)
    generate_word_report(df, anomalies, avg_value, chart, output_report)
